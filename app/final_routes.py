import json
import os
import re
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt
from html import unescape
import os
from urllib.parse import urlencode
from flask import Blueprint,  redirect, render_template,jsonify, request, send_file, url_for
from flask_cors import CORS # type: ignore
import numpy as np
import requests 
from .extensions import db
from sqlalchemy import text
from functools import wraps
from .configuration import configure
from dotenv import load_dotenv
from collections import defaultdict
from flask import session
from bs4 import BeautifulSoup
import os
import pickle
import json
from collections import defaultdict


main = Blueprint('main',__name__)

def rename_folder(old_path, new_name):
    parent_dir = os.path.dirname(old_path)
    new_path = os.path.join(parent_dir, new_name)
    os.rename(old_path, new_path)
    print(f"Renamed folder: {old_path} -> {new_path}")
    return new_path

# Function to rename file
def rename_file(folder_path, old_name, new_name):
    old_file_path = os.path.join(folder_path, old_name + ".docx")
    new_file_path = os.path.join(folder_path, new_name + ".docx")
    os.rename(old_file_path, new_file_path)
    print(f"Renamed file: {old_file_path} -> {new_file_path}")
    return new_file_path

def clean_text(raw_content):
    if raw_content!=None:
        """
        Cleans up impure XML content by removing <fo> tags, converting them into paragraphs,
        decoding entities, and fixing whitespace while preserving paragraph structure.
        """
        # Parse the raw content using BeautifulSoup
        soup = BeautifulSoup(raw_content, "html.parser")
    
        # Replace all <fo:block> and <fo:inline> tags with <p> tags
        for fo_tag in soup.find_all(['fo:block', 'fo:inline']):
            fo_tag.name = 'p'
    
        # Gather all paragraphs
        paragraphs = []
        for paragraph in soup.find_all('p'):
            # Extract text, clean up extra whitespace, and append to the list
            paragraph_text = paragraph.get_text(strip=True)
            if paragraph_text:  # Ignore empty paragraphs
                paragraphs.append(paragraph_text)
    
        # Decode HTML entities and join paragraphs with double newlines
        cleaned_text = "\n\n".join(unescape(paragraph) for paragraph in paragraphs)
    
        return cleaned_text
    else:
        return ""

def calculate_target_word_count(current_word_count):
    """Calculate the target word count by reducing 10%."""
    return round(current_word_count * 0.9)

def create_word_doc(folder_path, filename, content, current_word_count, data,image_path):
    """
    Create a Word document with the specified content, word counts, and optional tabular data.
    
    :param folder_path: Path to save the Word document.
    :param filename: Name of the Word document.
    :param content: Text content to add to the document.
    :param current_word_count: Current word count of the content.
    :param tabular_data: Optional tabular data as a list of dictionaries. Each dictionary represents a row.
    """
    target_word_count = calculate_target_word_count(current_word_count)
    doc = Document()

    # Add Current and Target word count
    paragraph = doc.add_paragraph()
    paragraph.alignment = 2  # Right-align
    paragraph.add_run(f"Current word count: {current_word_count}\n").bold = True
    paragraph.add_run(f"Target word count: {target_word_count}\n").bold = True

    # Add checklist
    checklist = [
        "AUTHOR",
        "EDITORIAL BOARD REVIEW",
        "AUTHOR 2ND SUBMISSION (IF NECESSARY)",
        "EDITORIAL BOARD 2ND REVIEW (IF NECESSARY)",
        "COPY EDITOR",
        "CBIS"
    ]
    for item in checklist:
        doc.add_paragraph(item)
        
    tabular_data = json.loads(data)
    # Add tabular data if provided and not empty
    if tabular_data and len(tabular_data) > 0:
        # Add a table with headers based on keys of the first dictionary
        headers = tabular_data[0].keys()
        table = doc.add_table(rows=0, cols=len(headers))
        table.style = 'Table Grid'
        
        grouped_data = defaultdict(list)
        for row in tabular_data:
            grouped_data[row['title']].append(row['candidateName'])
        
        # Prepare data for the table
        processed_data = [{'title': title, 'candidateNames': ', '.join(candidates)}
                        for title, candidates in grouped_data.items()]

        for row_data in processed_data:
            row_cells = table.add_row().cells
            row_cells[0].text = row_data['title']
            row_cells[1].text = row_data['candidateNames']
        # Add profile content
        
    doc.add_paragraph("")
    
    if image_path and os.path.exists(image_path):
        doc.add_paragraph()  # Add some space before the image
        doc.add_picture(image_path, width=Pt(300),height=Pt(300))
    
    doc.add_paragraph(content)
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)
    
    # Save the document
    doc_path = os.path.join(folder_path, filename + ".docx")
    doc.save(doc_path)
    print(f"Created: {doc_path}")

def congressional_dist(state_id,folder_path, doc_name,img_path):
    sql_tab = text('''
                select m.title,
                ISNULL((SELECT CAST(d.value AS VARCHAR(5)) + 'D  '  
                FROM ALMC_StateCongress_Details d WHERE d.SCong_ID = m.SCong_ID 
                and d.PartyCode = 'D' AND ISNULL(d.value,0) <> 0),'') + 
                ISNULL((SELECT CAST(d.value AS VARCHAR(5)) + 'R  ' 
                FROM ALMC_StateCongress_Details d where d.SCong_ID = m.SCong_ID and d.PartyCode = 'R'  AND ISNULL(d.value,0) <> 0),'') + 
                ISNULL((SELECT CAST(d.value AS VARCHAR(5)) + 'V'  FROM ALMC_StateCongress_Details d WHERE d.SCong_ID = m.SCong_ID and d.PartyCode = 'V' AND ISNULL(d.value,0) <> 0),'') Lineup
                from ALMC_StateCongress_Master m
                    inner join almc_state s on s.state_id = m.state_id
                WHERE m.State_ID = :state_id
                Order by s.state_name, m.displayorder
            ''')
    with db.engine.connect() as conn:
        results_tab = conn.execute(sql_tab,{'state_id':state_id}).fetchall()
        conn.commit()
        conn.close()
    tab_data= []
    for row in results_tab:
        tab_data.append({'title':row[0],'candidateName':row[1]})
    
    sql = text('''
                select ISNULL(State_CongressionalDists,'')AS State_CongressionalDists 
                from almc_state WHERE State_ID = :state_id
            ''')
    with db.engine.connect() as conn:
        results = conn.execute(sql,{'state_id':state_id}).fetchall()
        conn.commit()
        conn.close()
    con_dist_content= []
    for row in results:
        con_dist_content.append({'content':row[0]})
    try:
        content = clean_text(con_dist_content[0]['content'])
        word_count = len(content.split())
        create_word_doc(folder_path, doc_name, content, word_count,json.dumps(tab_data),img_path)
    except Exception as e:
        print(e)
##

def president_politics(state_id,folder_path, doc_name,img_path):
    sql_tab = text('''
                SELECT Title,CandidateName,TotalVotes,VotePerc 
                from ALMC_StatePresidentVote
                WHERE state_ID = :state_id AND election = 'Primary'
                    AND  cast(replace(VotePerc,'%','') as decimal) >= 5
                ORDER BY ElectedYear Desc,title, cast(replace(VotePerc,'%','') as decimal) Desc
            ''')
    with db.engine.connect() as conn:
        results_tab = conn.execute(sql_tab,{'state_id':state_id}).fetchall()
        conn.commit()
        conn.close()
    tab_data= []
    if results_tab:
        for row in results_tab:
            tab_data.append({'title':row[0],'candidateName':row[1]})
    else:
        pass
    
    sql = text('''
                select ISNULL(State_PresidentialPolitics,'')as State_PresidentialPolitics
	from almc_state WHERE State_ID = :state_id
            ''')
    with db.engine.connect() as conn:
        results = conn.execute(sql,{'state_id':state_id}).fetchall()
        conn.commit()
        conn.close()
    con_dist_content= []
    if results:
        for row in results:
            con_dist_content.append({'content':row[0]})
        try:
            content = clean_text(con_dist_content[0]['content'])
            word_count = len(content.split())

            create_word_doc(folder_path, doc_name, content, word_count,json.dumps(tab_data),img_path)
        except Exception as e:
            print(e)
    else:
        pass
##

def gov_profile(state_id,folder_path, doc_name,img_path):
    sql = text('''
                select GovernorName, Governor_image,Governor_writeup from ALMC_Governor where state_id = :state_id
            ''')
    with db.engine.connect() as conn:
        results = conn.execute(sql,{'state_id':state_id}).fetchall()
        conn.commit()
        conn.close()
    con_dist_content= []
    if results:
        for row in results:
            con_dist_content.append({'name':row[0],'content':row[2]})
        
        try:
            content = clean_text(con_dist_content[0]['content'])
            word_count = len(content.split())
            
            new_name =  doc_name.replace('Placeholder Governor',con_dist_content[0]['name'])
            
            create_word_doc(folder_path, doc_name, content, word_count,"[]",img_path)
            
            rename_file(folder_path,doc_name,new_name)
        except Exception as e:
            print(e)
    else:
        pass
##

def state_profile(state_id,folder_path, doc_name,img_path):
    sql = text('''
                select ISNULL(State_writeUP,'') as State_writeUP from almc_state WHERE State_ID = :state_id
            ''')
    with db.engine.connect() as conn:
        results = conn.execute(sql,{'state_id':state_id}).fetchall()
        conn.commit()
        conn.close()
    if results:
        con_dist_content= []
        for row in results:
            con_dist_content.append({'content':row[0]})
        try:
            content = clean_text(con_dist_content[0]['content'])
            word_count = len(content.split())

            create_word_doc(folder_path, doc_name, content, word_count,"[]",img_path)
        except Exception as e:
            print(e)
    else:
        pass
        
##

def map_fun(state_id,folder_path, doc_name,img_path):
    print(img_path)
    sql = text('''
                select State_Map from almc_state where state_id = :state_id
            ''')
    with db.engine.connect() as conn:
        results = conn.execute(sql,{'state_id':state_id}).fetchall()
        conn.commit()
        conn.close()
    con_dist_content= []
    if results:
        for row in results:
            con_dist_content.append({'content':row[0]})
        try:
            content = clean_text(con_dist_content[0]['content'])
            word_count = len(content.split())

            create_word_doc(folder_path, doc_name, content, word_count,"[]",img_path)
        except Exception as e:
            print(e)
    else:
        pass
##

 
@main.route('/',methods=["GET"])
def retrive_state():
    sql = text('''select State_id,State_name,State_Map,State_code, is_territory from almc_state order by state_name ''')
    with db.session.begin():
        result = db.session.execute(sql).fetchall()
    
    state_info =[]
    for row in result:
        state_info.append({
            'state_id':row[0],
            'state_name':row[1],
            'state_map':row[2],
            'state_code':row[3]
        })
    for state in state_info:
        state_name = state["state_name"]
        state_acronym = state["state_code"]
        base_path = r'D:\almanac_2026\almanac-2026'

        # Create main state folder
        state_folder = os.path.join(base_path, f"2026 - {state_name}")
        os.makedirs(state_folder, exist_ok=True)

        # Define subfolders and files
        folders = {
            "Congressional Districts": f"2026 - Cong. Districts – {state_acronym}",
            "District Profiles": f"2026 - Dist. Profile – {state_acronym}",
            "Gov. Profile": f"2026 - Gov. Profile – {state_acronym} - Placeholder Governor",
            "Map": f"2026 – State Map – {state_acronym}",
            "Presidential Politics": f"2026 – Presidential Politics – {state_acronym}",
            "State Profile": f"2026 - State Profile – {state_acronym}"
        }

        for folder, doc_name in folders.items():
            folder_path = os.path.join(state_folder, folder)
            os.makedirs(folder_path, exist_ok=True)
                
            if folder== "Congressional Districts":
                congressional_dist(state["state_id"],folder_path, doc_name,"")
            elif folder== "Presidential Politics":
                president_politics(state["state_id"],folder_path, doc_name,"")
            elif folder== "Gov. Profile":
                gov_profile(state["state_id"],folder_path, doc_name,"")
            elif folder== "State Profile":
                state_profile(state["state_id"],folder_path, doc_name,"")
            elif folder == "Map":
                img_path =  r"D:\almanac_2026\State Images"
                map_fun(state["state_id"],folder_path, doc_name,img_path +f"\{str(state["state_map"])}")

        # Additional folders like Representative and Senator Profiles
        rep_profiles_path = os.path.join(state_folder, "Rep. Profiles")
        os.makedirs(rep_profiles_path, exist_ok=True)

        # Placeholder Representative profiles
        
        sql = text('''  select HouseRepName,HouseRep_Writeup
                        from ALMC_State_HouseRep where state_id = :state_id order by DisplayOrder ''')
        with db.session.begin():
            result = db.session.execute(sql,{'state_id':state["state_id"]}).fetchall()
        
        rep_info =[]
        for row in result:
            rep_info.append({
                'name':row[0],
                'content':row[1],

            })
        count=1
        for row in rep_info: 
            rep_name = f"{row['name']}"
            folder_name = f"Rep. {count} – {rep_name}"
            rep_folder = os.path.join(rep_profiles_path, folder_name)
            os.makedirs(rep_folder, exist_ok=True)

            content = clean_text(row['content'])
            word_count = len(content.split())
            doc_name = f"2026 - Sen. {count} – {state_acronym} – {rep_name}"
            count+=1
            create_word_doc(rep_folder, doc_name, content, word_count,"[]","")

        sep_profiles_path = os.path.join(state_folder, "Sen. Profiles")
        os.makedirs(sep_profiles_path, exist_ok=True)
        
        sql = text('''  select case when Title = 'Senior Senator' then 1 else 2 end DisplayOrder,SenatorName, Senator_Writeup
                        from almc_State_senator where state_id = :state_id order by DisplayOrder ''')
        with db.session.begin():
            result = db.session.execute(sql,{'state_id':state["state_id"]}).fetchall()
        
        rep_info =[]
        for row in result:
            rep_info.append({
                'name':row[1],
                'content':row[2],

            })
            
        count=1
        for row in rep_info:  
            rep_name = f"{row['name']}"
            folder_name = f"Sen. {count} – {rep_name}"
            sen_folder = os.path.join(sep_profiles_path, folder_name)
            os.makedirs(sen_folder, exist_ok=True)

            content = clean_text(row['content'])
            word_count = len(content.split())
            doc_name = f"2026 - Sen {count} – {state_acronym} – {rep_name}"
            count+=1
            create_word_doc(sen_folder, doc_name, content, word_count,"[]","")
            
        dist_profiles_path = os.path.join(state_folder, "District Profiles")
        os.makedirs(dist_profiles_path, exist_ok=True)
        
        sql = text('''select DisplayOrder,DistrictName, Isnull(DistrictHeader + ': ','') + Isnull(District_Writeup,'') District_Writeup
                        from ALMC_State_HouseRep where state_id = :state_id order by DisplayOrder ''')
        with db.session.begin():
            result = db.session.execute(sql,{'state_id':state["state_id"]}).fetchall()
        
        dist_info =[]
        for row in result:
            dist_info.append({
                'name':row[1],
                'content':row[2],

            })
        count=1
        for row in dist_info:  
            rep_name = "District"
            folder_name = f"{count} – {rep_name}"
            dist_folder = os.path.join(dist_profiles_path, folder_name)
            os.makedirs(dist_folder, exist_ok=True)

            content = clean_text(row['content'])
            word_count = len(content.split())
            doc_name = f"2026 - Dist. {count} – {state_acronym}"
            count+=1
            create_word_doc(dist_folder, doc_name, content, word_count,"[]","")
            
        print(f"Folder structure and documents for {state_name} created successfully.")
        
    return jsonify(state_info)